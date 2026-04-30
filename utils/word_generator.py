import logging
import os
import re
import glob
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont
from decimal import Decimal, ROUND_UP
from playwright.sync_api import Page, expect

class WordGenerator:
    def __init__(self, page: Page):
        self.page = page

    def start_document(self):
        report_title = f"Report".upper()
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)

        now = datetime.now()
        start = now.strftime("%a %B %d, %Y %H:%M:%S")

        p = doc.add_heading(report_title, 0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.bold = True
        doc.add_heading(report_title, 1)
        p = doc.add_paragraph()
        p.add_run(f"Generated on: {start}").italic = True

        return doc

    def add_screenshot_with_description(self, doc, text, image):
        path = 'screenshots/'
        now = datetime.now()
        date = now.strftime("%d %m %Y, %Y %H:%M:%S")
        p = doc.add_paragraph()
        p.add_run(text)
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True
        self.page.screenshot(path=f"{path} {image}", full_page=True)
        imgdraw = Image.open(f"{path} {image}")
        drawing = ImageDraw.Draw(imgdraw)
        drawing.text((10, imgdraw.height -25), date, fill=(255, 0, 0))
        imgdraw.save(f"{path} {image}")
        doc.add_picture(f"{path} {image}", width=Cm(19))
        return doc

    def add_screenshot_only(self, doc, image):
        path = 'screenshots/'
        self.page.screenshot(path=f"{path} {image}", full_page=True)
        doc.add_picture(f"{path} {image}", width=Cm(19))
        return doc

    def add_heading(self, doc, text, level):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.bold = True
        return doc

    def add_heading_fail(self, doc, text, level):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.bold = True
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)
        return doc

    def add_text(self, doc, text):
        p = doc.add_paragraph(text)
        return doc

    @staticmethod
    def clear_temp_screenshot(path):
        for file in glob.glob(f"{path}*.png"):
            os.remove(file)
        logging.info("Temporary screenshots cleared.")

    @staticmethod
    def clear_temp_words(path):
        for file in glob.glob(f"{path}*.docx"):
            os.remove(file)
        logging.info("Temporary Word documents cleared.")



